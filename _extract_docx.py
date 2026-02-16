import os
from docx import Document

BASE = r"c:\Users\baike\OneDrive\Documentos\Alcaldia\Tatiana\5-feb-2026\informe"
FILES = [
    "ACTA PARCIAL.docx",
    "INF. GESTION.docx",
    "INF.SUPERVISION.docx",
]

def extract(docx_path: str) -> list[str]:
    doc = Document(docx_path)
    lines: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    lines.append(text)
    return lines

fields_per_file: dict[str, set[str]] = {}
text_per_file: dict[str, list[str]] = {}

for name in FILES:
    path = os.path.join(BASE, name)
    lines = extract(path)
    text_per_file[name] = lines
    fields: set[str] = set()
    for line in lines:
        if ":" in line:
            left = line.split(":", 1)[0].strip()
            if left and len(left) <= 80:
                fields.add(left)
    fields_per_file[name] = fields

common_fields = set.intersection(*fields_per_file.values()) if fields_per_file else set()

print("FILES:")
for name in FILES:
    print("-", name, "lines", len(text_per_file[name]), "fields", len(fields_per_file[name]))

print("\nCOMMON_FIELDS:")
for field in sorted(common_fields):
    print("-", field)

for name, lines in text_per_file.items():
    out_path = os.path.join(BASE, name + ".txt")
    with open(out_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))

print("\nTXT_EXPORT_DONE")
