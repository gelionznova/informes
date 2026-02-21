import os
import re
from docx import Document

ROOT = os.path.dirname(os.path.abspath(__file__))
TEMPLATES_DIR = os.path.join(ROOT, "docx_templates")

EXPECTED = {
    "informe_no",
    "contrato_no",
    "fecha_contrato",
    "objeto_contractual",
    "entidad_contratante",
    "nit",
    "contratista",
    "cc",
    "ciudad_expedicion",
    "plazo_dia",
    "plazo_mes",
    "plazo_anio",
    "fecha_inicio_contrato",
    "fecha_vencimiento_contrato",
    "valor_inicial",
    "cdp",
    "rp",
    "periodo_i_de",
    "periodo_i_a",
    "fecha_expedicion_informe_dia",
    "fecha_expedicion_informe_mes",
    "fecha_expedicion_informe_año",
    "fecha_presentacion_informe",
    "obligaciones_directas",
    "obligaciones_directas_ejecutadas",
    "obligaciones_generales",
    "obligaciones_generales_ejecutadas",
    "aportes_planilla",
    "aportes_mes",
    "fecha_pago_aportes",
    "operador_planilla",
    "aportes_valor_salud",
    "aportes_valor_pension",
    "aportes_valor_riesgos",
    "aportes_valor_caja_compensacion_familiar",
    "total_aportes",
    "valor_contrato",
    "valor_anticipo",
    "valor_pago_anticipado",
    "valor_adiciones",
    "valor_ejecutado",
    "valor_a_cobrar",
    "saldo_pendiente",
}

PLACEHOLDER_RE = re.compile(r"\{\{\s*([\w\u00f1\u00d1]+)\s*\}\}")


def extract_placeholders(doc_path: str) -> set[str]:
    doc = Document(doc_path)
    text_chunks = []
    for paragraph in doc.paragraphs:
        text_chunks.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text_chunks.append(cell.text)
    placeholders = set()
    for chunk in text_chunks:
        for match in PLACEHOLDER_RE.findall(chunk):
            placeholders.add(match)
    return placeholders


def main() -> None:
    if not os.path.isdir(TEMPLATES_DIR):
        print("TEMPLATES_DIR_NOT_FOUND")
        return

    files = [f for f in os.listdir(TEMPLATES_DIR) if f.lower().endswith(".docx")]
    if not files:
        print("NO_TEMPLATES_FOUND")
        return

    for filename in sorted(files):
        path = os.path.join(TEMPLATES_DIR, filename)
        found = extract_placeholders(path)
        missing = sorted(EXPECTED - found)
        extra = sorted(found - EXPECTED)
        print(f"TEMPLATE: {filename}")
        print(f"FOUND: {len(found)}")
        print("MISSING:")
        for item in missing:
            print(f"- {item}")
        print("EXTRA:")
        for item in extra:
            print(f"- {item}")
        print("---")


if __name__ == "__main__":
    main()
