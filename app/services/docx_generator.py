import base64
import io
import math
import os
from docxtpl import DocxTemplate
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
ROOT_DIR = os.path.dirname(BASE_DIR)
TEMPLATES_DIR = os.path.join(ROOT_DIR, "docx_templates")
OUTPUT_DIR = os.path.join(ROOT_DIR, "output")

TEMPLATE_FILES = {
    "inf_gestion": "INF. GESTION.docx",
    "inf_supervision": "INF.SUPERVISION.docx",
    "acta_parcial": "ACTA.docx",
    "anexo": "ANEXO.docx",
}

def _normalize(text: str) -> str:
    return " ".join(text.upper().split())

def _cell_has_keywords(text: str, keywords: list[str]) -> bool:
    normalized = _normalize(text)
    return all(keyword in normalized for keyword in keywords)

def _find_gestion_table(doc: Document):
    for table in doc.tables:
        for row_index, row in enumerate(table.rows):
            if len(row.cells) < 2:
                continue
            if _cell_has_keywords(row.cells[0].text, ["ACTIVIDADES", "CONTRATO"]) and _cell_has_keywords(
                row.cells[1].text, ["EJECUT"]
            ):
                return table, row_index
    return None, None

def _find_supervision_table(doc: Document):
    for table in doc.tables:
        for row_index, row in enumerate(table.rows):
            if len(row.cells) < 3:
                continue
            matches_contrato = _cell_has_keywords(
                row.cells[0].text, ["ACTIVIDADES", "CONTRATO"]
            ) and _cell_has_keywords(row.cells[1].text, ["EJECUT"])
            matches_reportadas = _cell_has_keywords(
                row.cells[0].text, ["ACTIVIDADES", "REPORT"]
            ) and _cell_has_keywords(row.cells[1].text, ["OBSERV"])
            if (matches_contrato or matches_reportadas) and _cell_has_keywords(
                row.cells[2].text, ["EVIDEN"]
            ):
                return table, row_index
    return None, None

def _clear_rows_after(table, keep_rows: int) -> None:
    # Remove all rows after keep_rows (0-based count)
    while len(table.rows) > keep_rows:
        table._tbl.remove(table.rows[-1]._tr)

def _fill_gestion_table(doc: Document, items: list[dict]) -> None:
    table, header_row_index = _find_gestion_table(doc)
    if table is None:
        return
    _clear_rows_after(table, header_row_index + 1)
    for idx, item in enumerate(items, start=1):
        row = table.add_row()
        left = f"{idx}. {item.get('actividad_contrato', '').strip()}".strip()
        right = item.get("actividad_ejecutada", "").strip()
        row.cells[0].text = left
        row.cells[1].text = right

def _fill_supervision_table(doc: Document, items: list[dict]) -> None:
    table, header_row_index = _find_supervision_table(doc)
    if table is None:
        return
    _clear_rows_after(table, header_row_index + 1)
    for idx, item in enumerate(items, start=1):
        row = table.add_row()
        left = f"{idx}. {item.get('actividad_contrato', '').strip()}".strip()
        right = item.get("actividad_ejecutada_tercera", "").strip()
        if not right:
            right = item.get("actividad_ejecutada", "").strip()
        evidencias = item.get("aporta_evidencias", "").strip()
        row.cells[0].text = left
        row.cells[1].text = right
        row.cells[2].text = evidencias

def _decode_data_url(value: str) -> bytes | None:
    if not value:
        return None
    if value.startswith("data:"):
        try:
            _, encoded = value.split(",", 1)
        except ValueError:
            return None
    else:
        encoded = value
    try:
        return base64.b64decode(encoded)
    except (ValueError, base64.binascii.Error):
        return None

def _build_anexo_document(context: dict, out_path: str) -> None:
    doc = Document()
    doc.add_heading("ANEXO", level=1)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    items = context.get("obligaciones_directas_items")
    if not isinstance(items, list):
        items = []
    if not items:
        doc.add_paragraph("No se registraron actividades con evidencias.")
        doc.save(out_path)
        return

    for idx, item in enumerate(items, start=1):
        actividad = item.get("actividad_contrato", "").strip()
        ejecutada = item.get("actividad_ejecutada", "").strip()
        paragraph = doc.add_paragraph(
            f"{idx}. Actividad del contrato: {actividad}"
        )
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph = doc.add_paragraph(f"Actividad ejecutada: {ejecutada}")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        evidencias = item.get("evidencias") or {}
        images = evidencias.get("images") or []
        groups = evidencias.get("groups") or []
        pdfs = evidencias.get("pdfs") or []
        total_groups = max(len(groups), math.ceil(len(images) / 3))

        if not images and not pdfs:
            doc.add_paragraph("Evidencias: Sin evidencia.")
            doc.add_paragraph("")
            continue

        if images:
            for group_index in range(total_groups):
                start = group_index * 3
                chunk = images[start : start + 3]
                if not chunk:
                    continue
                table = doc.add_table(rows=1, cols=3)
                for cell_index, image in enumerate(chunk):
                    data_url = image.get("dataUrl") or image.get("data_url")
                    image_bytes = _decode_data_url(data_url)
                    if not image_bytes:
                        continue
                    run = table.rows[0].cells[cell_index].paragraphs[0].add_run()
                    # Tamaño fijo: 4x4 cm por imagen
                    from docx.shared import Cm

                    run.add_picture(
                        io.BytesIO(image_bytes), width=Cm(4), height=Cm(4)
                    )

                group = groups[group_index] if group_index < len(groups) else {}
                description = (group.get("description") or "").strip()
                date = (group.get("date") or "").strip()
                paragraph = doc.add_paragraph(
                    f"Descripcion: {description}" if description else "Descripcion:"
                )
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                paragraph = doc.add_paragraph(
                    f"Fecha: {date}" if date else "Fecha:"
                )
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                doc.add_paragraph("")
        else:
            doc.add_paragraph("Evidencia fotografica: Sin evidencia.")

        if pdfs:
            paragraph = doc.add_paragraph("Evidencia PDF:")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for pdf_index, pdf in enumerate(pdfs, start=1):
                pdf_name = ""
                if isinstance(pdf, dict):
                    pdf_name = str(pdf.get("name", "")).strip()
                if not pdf_name:
                    pdf_name = f"Documento {pdf_index}.pdf"
                paragraph = doc.add_paragraph(f"- {pdf_name}")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            paragraph = doc.add_paragraph("Evidencia PDF: Sin evidencia.")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        doc.add_paragraph("")

    doc.save(out_path)

def generate_documents(context: dict, record_id: int) -> dict:
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    outputs = {}
    for key, filename in TEMPLATE_FILES.items():
        if key == "anexo":
            out_name = f"{record_id:05d}_{filename}"
            out_path = os.path.join(OUTPUT_DIR, out_name)
            _build_anexo_document(context, out_path)
            outputs[key] = out_path
            continue
        template_path = os.path.join(TEMPLATES_DIR, filename)
        doc = DocxTemplate(template_path)
        doc.render(context)
        out_name = f"{record_id:05d}_{filename}"
        out_path = os.path.join(OUTPUT_DIR, out_name)
        temp_path = out_path + ".tmp"
        doc.save(temp_path)

        if key in {"inf_gestion", "inf_supervision"}:
            items = context.get("obligaciones_directas_items")
            if isinstance(items, list):
                word_doc = Document(temp_path)
                if key == "inf_gestion":
                    _fill_gestion_table(word_doc, items)
                else:
                    items_supervision = context.get(
                        "obligaciones_directas_items_tercera", items
                    )
                    if not isinstance(items_supervision, list):
                        items_supervision = items
                    _fill_supervision_table(word_doc, items_supervision)
                word_doc.save(out_path)
                try:
                    os.remove(temp_path)
                except OSError:
                    pass
            else:
                os.replace(temp_path, out_path)
        else:
            os.replace(temp_path, out_path)
        outputs[key] = out_path
    return outputs
